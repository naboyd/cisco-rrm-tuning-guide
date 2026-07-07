#!/usr/bin/env python3
"""Generate RRM-Settings-Source-References.docx from validated Cisco sources."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "RRM-Settings-Source-References.docx"

PRIMARY_SOURCES = [
    (
        "C9800 RRM Configuration Guide (17.12.x)",
        "https://www.cisco.com/c/en/us/td/docs/wireless/controller/9800/17-12/config-guide/b_wl_17_12_cg/m_rrm_c9800.html",
        "Factory defaults, DCA/TPC/CHD/ED-RRM CLI, per-band tpc-threshold, Avoid Foreign AP default enabled.",
    ),
    (
        "CX Design Guide — Wireless for Large Public Networks",
        "https://www.cisco.com/c/en/us/support/docs/wireless/catalyst-9800-series-wireless-controllers/222000-design-guide-cx-wireless-for-large-pub.html",
        "Stadium/arena HD: CHD off, ED-RRM off, directional antenna power from survey, RxSOP, 20 MHz, FRA not recommended.",
    ),
    (
        "C9800 Configuration Best Practices",
        "https://www.cisco.com/c/en/us/td/docs/wireless/controller/9800/technical-reference/c9800-best-practices.html",
        "Site tags for WNCD load (not RRM); stadium 3000 AP / 8 tags example; custom site tags vs default.",
    ),
    (
        "C9800 RRM Deployment Guide (Technote)",
        "https://www.cisco.com/c/en/us/td/docs/wireless/controller/technotes/8-8/b_C9800_rrm_dg.html",
        "DCA anchor time 4 hr interval for voice; channel width-max 40 MHz guidance.",
    ),
    (
        "RRM White Paper — Coverage Hole Detection",
        "https://www.cisco.com/c/en/us/td/docs/wireless/controller/technotes/8-3/b_RRM_White_Paper/chd.pdf",
        "CHD algorithm behavior, min failed clients, power increase mitigation.",
    ),
    (
        "C9800 RF Profile Configuration Guide (17.3.x)",
        "https://www.cisco.com/c/en/us/td/docs/wireless/controller/9800/17-3/config-guide/b_wl_17_3_cg/m_rf_profile_ewlc.html",
        "University HD areas, RxSOP, high-density RF profile parameters.",
    ),
    (
        "Catalyst Center AI-Enhanced RRM Deployment Guide",
        "https://www.cisco.com/c/en/us/td/docs/wireless/controller/9800/technical-reference/ai-enhanced-rrm-dg.html",
        "AI RF Profile services, busy hour, migration from classic RRM.",
    ),
    (
        "C9800 RF-based Automatic AP Load Balancing (17.18.x)",
        "https://www.cisco.com/c/en/us/td/docs/wireless/controller/9800/17-18/config-guide/b_wl_17_18_cg/m_auto-wncd-lb.html",
        "17.12+ alternative to manual site-tag WNCD balancing.",
    ),
    (
        "Troubleshoot WLC CPU Load (WNCD / site tags)",
        "https://www.cisco.com/c/en/us/support/docs/wireless/catalyst-9800-series-wireless-controllers/221965-troubleshoot-wireless-lan-controller-cpu.html",
        "Site tag assignment to WNCD; load command; inter-WNCD roaming considerations.",
    ),
    (
        "C9104 Stadium Antenna Deployment Guide",
        "https://www.cisco.com/c/en/us/td/docs/wireless/antennas/install/guide/c9104_deployment_guide.html",
        "Directional stadium antenna overlap; power selection caveats.",
    ),
]

SETTING_SOURCES = [
    ("DCA Interval", "10 min (0 = 10 min)", "4–8 hours + anchor time",
     "C9800 CG 17.12 RRM; CX Large Public Networks",
     "Default interval 10 minutes. CX: must increase from 10 min in all cases; example 4 hr + anchor 8.",
     "Supported"),
    ("DCA Sensitivity", "Medium (15 dB)", "Low in multi-tenant / HD foreign RF",
     "C9800 CG 17.12 RRM",
     "Medium = 15 dB default; Low = 30 dB less sensitive to environmental drift.",
     "Supported"),
    ("DCA Anchor Time", "Not set", "Off-peak hour (0–23)",
     "C9800 CG 17.12 RRM; C9800 RRM DG",
     "Anchortime pairs with long interval; RRM DG recommends 4 hr + anchor for voice environments.",
     "Supported"),
    ("Avoid Foreign AP", "Enabled (global & profile)", "Off in multi-tenant / noisy foreign RF",
     "C9800 CG 17.12 RRM; CX Large Public Networks",
     "Default enabled. CX: rogue/foreign avoidance can destabilize when many foreign BSSIDs present.",
     "Supported — context-dependent"),
    ("Avoid Cisco AP Load", "Disabled", "Keep disabled",
     "C9800 CG 17.12 RRM",
     "Explicit default disabled in configuration guide.",
     "Supported"),
    ("ED-RRM", "Enabled (with CleanAir)", "Off in busy HD / multi-tenant",
     "C9800 CG RRM/CleanAir; CX Large Public Networks",
     "ED-RRM triggers rapid channel change. CX: do not use in busy, noisy, high-density environments.",
     "Supported"),
    ("ED-RRM Rogue Contribution", "Configurable (often off by default)", "Off when neighbors trigger false events",
     "C9800 CleanAir CG",
     "Rogue duty-cycle channel changes; disable when neighbor WLANs classified as rogue.",
     "Supported"),
    ("CHD (Coverage Hole Detection)", "Enabled", "On enterprise; Off stadium/large public",
     "CX Large Public Networks; CHD White Paper; C9800 CG",
     "CX: CHD must be disabled globally for large public networks. CHD raises tx power per radio.",
     "Supported"),
    ("CHD Min Failed Clients", "3 clients", "5–8 in dense / foreign RF",
     "C9800 CG CHD; CHD White Paper",
     "Minimum clients before mitigation raises power; tuning reduces false positives.",
     "Interpretation — field tuning"),
    ("TPC Update Interval", "600 sec", "3600 sec (align with DCA)",
     "C9800 CG 17.12 RRM; CX Large Public",
     "Default 600 s TPC; CX notes TPC on 10 min schedule. Align TPC cadence with DCA philosophy.",
     "Supported"),
    ("TPC Min / Max", "Unbounded / max 30 dBm", "Set per RF design per band",
     "C9800 CG 17.12 RRM; CX Large Public Networks",
     "CX: TPC must be limited to appropriate range. Directional: narrow or fixed range from survey.",
     "Supported"),
    ("TPC Power Threshold", "-70 dBm per band", "From survey; configure per 24/5/6 GHz",
     "C9800 CG 17.12 RRM (ap dot11 {24|5|6}ghz rrm tpc-threshold); CX Large Public",
     "CLI is per band. CX: omni examples 5 dBm min; directional values are survey outcomes.",
     "Supported — per band required"),
    ("TPC Channel Aware (5 GHz)", "Varies", "Enabled in HD per CX",
     "CX Large Public Networks",
     "CX recommends enable TPC Channel Aware in high-density environments.",
     "Supported"),
    ("Channel Width 5 GHz", "20 MHz", "20 MHz dense; 40 MHz owned/low density",
     "CX Large Public Networks; C9800 RRM DG",
     "CX HD: 20 MHz. RRM DG: cap width-max 40 in many environments.",
     "Supported"),
    ("RxSOP", "Auto", "Custom dBm from survey (per band)",
     "CX Large Public Networks; C9800 RF Profile CG",
     "CX extensive RxSOP section for HD venues; defines receive cell boundary.",
     "Supported — design-dependent"),
    ("Data Rates (5 GHz mandatory)", "12 Mbps typical default", "12 Mbps campus; 18–24 Mbps HD",
     "CX Large Public Networks",
     "CX: 12 Mbps low density; 24 Mbps HD stadium/events with testing.",
     "Supported"),
    ("FRA (Flexible Radio Assignment)", "Often disabled", "Off large public; on XOR enterprise only",
     "CX Large Public Networks",
     "CX: FRA not recommended for large public; deterministic config preferred.",
     "Supported"),
    ("Client Network Preference", "default", "5GHz when 2.4 saturated",
     "C9800 RF Profile CG",
     "RF profile setting: default / 5ghz / 6ghz.",
     "Supported"),
    ("Client-Aware FRA client-select %", "50%", "50–60%+ before 2.4→5G in dense RF",
     "C9800 CG / show ap fra",
     "Guards FRA band conversion until enough clients on target band.",
     "Interpretation — field tuning"),
    ("Client-Aware FRA client-reset count", "5 clients", "5–10+ by density",
     "C9800 CG / FRA documentation",
     "Minimum clients before FRA band reset.",
     "Interpretation — field tuning"),
    ("Band Select (WLAN)", "Varies per WLAN", "On data WLANs; off voice",
     "Enterprise WLAN best practices",
     "WLAN-level; not part of RRM algorithm.",
     "Supported"),
    ("Site Tags (WNCD)", "default-site-tag", "Custom tags for controller scale — NOT RRM",
     "C9800 Best Practices; Config Model Limitations; CPU Load doc",
     "Site tags balance WNCD CPU/memory. Not an RF optimization setting. 17.12+ RF load balancing alternative.",
     "Corrected in guide v2 — architecture only"),
    ("Stadium DCA Freeze workflow", "N/A", "Engineered plan — do not auto-converge-then-freeze",
     "CX Large Public Networks; field review (Al)",
     "CX mentions freeze after stable auto DCA for some venues; directional stadium antennas require survey-led fixed plan. Guide corrected per SME feedback.",
     "Corrected — was over-applied"),
    ("AI-Enhanced RRM Busy Hour", "Admin configured", "Match site peak hours",
     "AI-Enhanced RRM DG; Catalyst Assurance UG 3.1",
     "Defers disruptive changes during busy hour; sensitivity Low for latency-sensitive sites.",
     "Supported"),
    ("AI-Enhanced RRM Services", "Legacy RRM default", "DCA/TPC/DBS/FRA per profile",
     "AI-Enhanced RRM DG",
     "At least one AI service required; WLC global RRM overridden for subscribed services.",
     "Supported"),
]

CORRECTIONS = [
    ("Stadium: auto DCA then freeze", "Removed from guide", "CX freeze guidance applies after a validated stable plan; directional C9104 deployments need survey-led channels/power, not empty-bowl RRM convergence."),
    ("Stadium: site tags in RRM table", "Moved to architecture docs only", "Site tags are WNCD scaling — C9800 Best Practices — not RRM parameters."),
    ("Single TPC threshold all bands", "Per-band callout added", "C9800 CG documents ap dot11 {24ghz|5ghz|6ghz} rrm tpc-threshold separately."),
    ("Archetype numbers without survey", "Disclaimer added", "CX: directional antenna min/max TPC typically from RF survey with narrow range."),
]


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("Cisco Wireless RRM Tuning Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Settings Source Reference & Validation Matrix")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Companion document to index.html (Cisco RRM Tuning Guide). "
        "Maps guide recommendations to published Cisco documentation. "
        "Generated for internal field use — not BU-validated."
    )
    doc.add_paragraph()

    add_heading(doc, "1. Purpose & Limitations", 1)
    doc.add_paragraph(
        "This document validates (or flags corrections for) settings in the RRM Tuning Guide against "
        "Cisco official configuration guides, design guides, and technotes. It does not replace RF "
        "survey, physical design, or BU support. Values marked 'Interpretation' reflect common field "
        "practice synthesized from Cisco guidance, not explicit single-value mandates."
    )
    doc.add_paragraph()

    add_heading(doc, "2. Primary Source Documents", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Document"
    hdr[1].text = "URL"
    hdr[2].text = "Relevance"
    for name, url, rel in PRIMARY_SOURCES:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = url
        row[2].text = rel
    doc.add_paragraph()

    add_heading(doc, "3. Setting-by-Setting Source Matrix", 1)
    doc.add_paragraph(
        "Guide profile values are starting points. 'C9800 Default' is from C9800 IOS-XE 17.12.x RRM "
        "configuration guide unless noted."
    )
    st = doc.add_table(rows=1, cols=6)
    st.style = "Table Grid"
    h = st.rows[0].cells
    for i, label in enumerate([
        "Setting", "C9800 Default", "Guide Direction",
        "Primary Cisco Source(s)", "Source Summary", "Status"
    ]):
        h[i].text = label
    for row_data in SETTING_SOURCES:
        cells = st.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    doc.add_paragraph()

    add_heading(doc, "4. Guide Corrections (SME Review)", 1)
    doc.add_paragraph(
        "Items corrected in the HTML guide following technical review (e.g. stadium directional "
        "antennas, site tag scope, per-band thresholds)."
    )
    ct = doc.add_table(rows=1, cols=3)
    ct.style = "Table Grid"
    ch = ct.rows[0].cells
    ch[0].text = "Original guide issue"
    ch[1].text = "Action"
    ch[2].text = "Rationale"
    for item in CORRECTIONS:
        r = ct.add_row().cells
        for i, val in enumerate(item):
            r[i].text = val
    doc.add_paragraph()

    add_heading(doc, "5. Deployment Profile Source Summary", 1)
    profiles = [
        ("Enterprise Campus", "C9800 RRM CG defaults + enterprise WLAN BP; Avoid Foreign AP may remain on in owned RF."),
        ("Multi-Tenant", "CX foreign/rogue guidance; disable Avoid Foreign AP, ED-RRM; 20 MHz."),
        ("NYC Downtown", "Ultra-dense interpretation of CX + multi-tenant themes; freeze optional after validated plan."),
        ("Atlanta Downtown", "Between campus and NYC; ED-RRM Low only if enabled."),
        ("Dorm / Residence", "C9800 RF Profile CG university HD; CX 20 MHz / RxSOP / ED-RRM off peak."),
        ("Stadium HD", "CX Large Public Networks primary; CHD/ED-RRM/FRA off; survey-led power/channels; C9104 DG for antennas."),
    ]
    pt = doc.add_table(rows=1, cols=2)
    pt.style = "Table Grid"
    pt.rows[0].cells[0].text = "Profile"
    pt.rows[0].cells[1].text = "Primary Cisco basis"
    for prof, basis in profiles:
        r = pt.add_row().cells
        r[0].text = prof
        r[1].text = basis
    doc.add_paragraph()

    add_heading(doc, "6. Community / Supplemental References", 1)
    doc.add_paragraph(
        "Cisco Community discussions are not official documentation but illustrate operational context:"
    )
    community = [
        ("WNCD balancing for ultra high-density (sports arena)", "https://community.cisco.com/t5/wireless/wncd-balancing-for-ultra-high-density/td-p/5260311"),
        ("DCA interval commissioning discussion", "https://community.cisco.com/t5/wireless/new-aps-cw9166-under-wlc9800-channel-not-auto-reconciled/td-p/5315264"),
        ("AI-Enhanced RRM deployment", "https://community.cisco.com/t5/cisco-catalyst-center/deploy-to-ai-enhanced-rrm/td-p/5291864"),
    ]
    for title_text, url in community:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{title_text}: ").bold = True
        p.add_run(url)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
